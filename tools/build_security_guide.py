from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "VITEC-Scoreboard-Deployment-and-Security-Guide.docx"
WEB = ROOT / "src" / "VS.Web" / "wwwroot" / "docs" / OUT.name

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 105, 120)
LIGHT = "E8EEF5"

def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None: tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")
    if tcW.getparent() is None: tcPr.append(tcW)

def table(rows, widths=None):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tPr = t._tbl.tblPr
    tblW = tPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    if tblW.getparent() is None: tPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tPr.append(tblInd)
    if widths is None: widths = [9360 // len(rows[0])] * len(rows[0])
    grid = t._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for ri, vals in enumerate(rows):
        cells = t.rows[0].cells if ri == 0 else t.add_row().cells
        for ci, val in enumerate(vals):
            set_cell_width(cells[ci], widths[ci]); cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val)); font(r, 9.5, bold=(ri == 0), color=DARK if ri == 0 else None)
            if ri == 0: shade(cells[ci], LIGHT)
    return t

def heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        a, b = text[:len(bold_prefix)], text[len(bold_prefix):]
        font(p.add_run(a), bold=True); font(p.add_run(b))
    else: font(p.add_run(text))
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.375); p.paragraph_format.first_line_indent = Inches(-.188)
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
        font(p.add_run(item))

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)
styles = doc.styles
normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK,10,5)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=True
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header = sec.header.paragraphs[0]; font(header.add_run("VITEC Scoreboard | Deployment and Security Guide"), 9, color=MUTED)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font(footer.add_run("Customer and administrator reference"), 8.5, color=MUTED)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(42); p.paragraph_format.space_after=Pt(8)
font(p.add_run("VITEC SCOREBOARD"), 26, bold=True, color=DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20)
font(p.add_run("Deployment, Network, Operations, and Cybersecurity Guide"), 15, color=BLUE)
table([["Document status","Applies to","Last updated"],["Operational reference","Windows 11; Windows Server 2022 and 2025","August 15, 2026"]],[2200,4200,2960])
para("Purpose. This guide identifies what VITEC Scoreboard installs, the services and processes it runs, the ports and protocols it uses, where it stores data, and the controls administrators should apply. It also defines a practical response process when a security vulnerability is reported.", "Purpose.")
para("Security position. VITEC Scoreboard is designed for a trusted internal network. The current web interface uses HTTP and does not provide application authentication or role-based access control. Do not publish it directly to the Internet or place it on an untrusted network segment.", "Security position.")

heading("1. Installed Components and Processes")
table([["Component","Windows identity / context","Purpose"],["VITEC Scoreboard service (VITECScoreboard)","LocalSystem","Hosts the web interface, API, MLB data retrieval, settings, and PostgreSQL integration."],["VITEC.Scoreboard.exe","Windows service process","Branded application host shown in Task Manager."],["VITEC.VideoOutput.exe","Signed-in interactive user","Embedded off-screen Chromium renderer and FFmpeg video-output coordinator."],["CefSharp subprocesses","Child processes of VITEC.VideoOutput.exe","Render the selected GameCenter view without using the customer’s installed browser or browser profile."],["FFmpeg","Child process started by the video helper","Encodes H.264 MPEG transport stream for UDP multicast or SRT output."],["PostgreSQL","Separate local or remote service","Stores normalized historical game and pitch records when configured."]],[2500,2400,4460])
para("The service currently runs as LocalSystem. Microsoft describes LocalSystem as highly privileged. Treat the executable, installer, settings directory, and update source as security-sensitive. Do not install this service on a domain controller.")

heading("2. Network Ports, Protocols, and Data Flows")
table([["Direction","Default endpoint","Protocol","Required for"],["Inbound","TCP 5000 (configurable)","HTTP","Scoreboard, GameCenter, Settings, JSON/XML feeds, and local output pages."],["Outbound","statsapi.mlb.com:443","HTTPS/TLS","MLB schedules, live feeds, statistics, and game data."],["Database","localhost:5432 (configurable)","PostgreSQL over TCP","Historical game and pitch storage."],["Outbound video","239.10.10.10:5004 (configurable)","UDP multicast / MPEG-TS / H.264","One-to-many video distribution on a managed network."],["Outbound video","Administrator-defined host and port","SRT caller over UDP","Point-to-point or routed low-latency video distribution."]],[1300,2250,1850,3960])
bullets(["Restrict inbound TCP access to approved management, signage, and EZ TV subnets. Do not use an Any/Internet scope unless the network security design explicitly requires it.","Keep Windows Defender Firewall enabled. Use Domain or Private profiles as appropriate and document every exception.","Multicast and SRT destinations must be approved by the network owner. Multicast is not encrypted or authenticated; control it with VLANs, routing, IGMP controls, and receiver access.","If remote PostgreSQL is used, restrict TCP 5432 to the application host and require PostgreSQL host-based access controls. Do not expose PostgreSQL to the Internet."])

heading("3. Files, Credentials, and Permissions")
table([["Location","Contents","Administrative guidance"],[r"C:\Program Files\VITEC\Scoreboard", "Installed executables, libraries, renderer, scripts, and static web content.","Administrators and SYSTEM should have write access; ordinary users should have read/execute only."],[r"C:\ProgramData\VITEC Scoreboard", "vssettings.json, logs, video settings, themes, uploaded advertising, and workspace templates.","Back up as configuration data. Restrict write access to administrators, SYSTEM, and the required application identities."],["PostgreSQL database", "Normalized game and pitch records and PostgreSQL accounts.","Use a unique 12–128 character password with uppercase, lowercase, number, and an approved special character; rotate when exposure is suspected."],[r"C:\ffmpeg\bin\ffmpeg.exe (default)", "Administrator-supplied FFmpeg executable.","Use an approved build, verify its origin/hash, and patch it with the rest of the application stack."]],[2600,3100,3660])
para("Credential warning. The PostgreSQL connection string, including the application password, is currently stored in the local settings file. Protect that file with NTFS permissions and backups appropriate for confidential configuration. A future hardening release should migrate the password to Windows protected storage.")

heading("4. Deployment Checklist")
bullets(["Use a supported, fully patched Windows 11, Windows Server 2022, or Windows Server 2025 system.","Install with an approved administrator account and verify the MSI and executable signatures when production signing is available.","Place the server on a trusted application or signage VLAN; never on a public-facing interface or domain controller.","Confirm the VITECScoreboard service is Automatic and running, and confirm its recovery actions restart it after failure.","Set the Webpage Port, then restart the service and update the firewall rule if the port changes.","Restrict inbound firewall scope to approved subnets. Validate outbound HTTPS to statsapi.mlb.com.","Provision the vsapp PostgreSQL account with a unique password. Do not reuse the PostgreSQL administrator password.","Verify JSON/XML feeds, GameCenter, database status, video preview, UDP multicast/SRT output, and log creation.","Record the installed VITEC Scoreboard, .NET, CefSharp/Chromium, FFmpeg, PostgreSQL, and Npgsql versions in the customer asset inventory.","Back up configuration and database data before upgrades; retain a tested rollback installer."])

heading("5. Current Security Limitations and Compensating Controls")
table([["Current limitation","Compensating control"],["HTTP only; no built-in TLS","Use only on a trusted segmented network. If remote access is required, place an approved authenticated TLS reverse proxy in front of the application."],["No application login or role-based access control","Restrict port 5000 (or the configured port) by firewall/subnet. Limit access to the Settings page to authorized administrator workstations."],["Service runs as LocalSystem","Protect the installer and install directory, monitor service changes, and do not deploy on a domain controller. Evaluate a lower-privilege service identity in a future release."],["Configuration contains a database credential","Apply restrictive NTFS ACLs, limit backup access, rotate on suspected exposure, and plan migration to DPAPI or another Windows secret store."],["UDP multicast is unencrypted","Use trusted VLANs, multicast boundaries, IGMP controls, and approved receivers. Use an appropriately secured SRT design when content must cross routed or untrusted links."],["Uploaded image/video media is processed and served locally","Permit uploads only from trusted administrators. Scan media with endpoint protection and retain the documented 250 MB application limit."],["Production signing may not yet be established","Distribute through an approved internal channel, record hashes, retain release provenance, and add Authenticode/MSI signing before broad customer deployment."]],[3300,6060])

heading("6. Vulnerability Management")
para("Monitor the complete software bill of materials, not only the VITEC application. At minimum, track VITEC Scoreboard, .NET, ASP.NET Core, NuGet dependencies, CefSharp and its Chromium runtime, FFmpeg, PostgreSQL, Npgsql, WiX installer tooling, and the Windows operating system.")
table([["Activity","Recommended cadence"],["Windows, Defender, and supported Microsoft product updates","Follow organizational patch policy; expedite actively exploited or critical vulnerabilities."],["PostgreSQL security page and supported minor releases","Review at every PostgreSQL security release and at least monthly."],[".NET and NuGet security advisories","Review with each monthly Microsoft security release and before every VITEC release."],["Chromium/CefSharp and FFmpeg security updates","Review at least monthly and before every VITEC release."],["Dependency/SBOM scan","Every build and release candidate; retain results with release records."],["External attack-surface and firewall review","At installation, after network changes, and at least annually."]],[4300,5060])
bullets(["Create and retain an SBOM for every released installer.","Record each finding with affected versions, exposure, severity, owner, mitigation, target release, and verification evidence.","Do not rely on version age alone. Evaluate reachability: exposed web routes, file upload paths, renderer content, network listeners, service privilege, and database permissions.","Test security updates in a representative Windows 11 and Windows Server environment before customer rollout."])

heading("7. Security Incident Response")
table([["Phase","Required actions"],["Identify","Capture product version, host, service/process state, logs, firewall rules, installed dependency versions, indicators, and the time the issue began."],["Contain","If active compromise is suspected, isolate the host or restrict its firewall access. Stop video output and the VITECScoreboard service if needed. Preserve logs and configuration before changing state."],["Assess","Determine affected versions, exploit prerequisites, data/credential exposure, lateral-movement risk, and whether LocalSystem or PostgreSQL privileges were reached."],["Remediate","Patch or upgrade affected components, rotate PostgreSQL and related credentials, remove unauthorized files/accounts/rules, and restore from trusted media where required."],["Validate","Re-scan, verify hashes/signatures, test service/API/database/video functions, review logs, and confirm firewall and account controls."],["Communicate","Notify affected customers through the approved security channel with impact, affected versions, mitigations, fixed version, and required actions. Avoid publishing exploit-enabling detail before containment."]],[1500,7860])

heading("8. Administrator Verification Commands")
table([["Purpose","PowerShell command"],["Service status",r"Get-Service VITECScoreboard"],["Service identity",r"Get-CimInstance Win32_Service -Filter \"Name='VITECScoreboard'\" | Select Name,StartName,State,PathName"],["Listening port",r"Get-NetTCPConnection -State Listen | Where-Object LocalPort -eq 5000"],["Firewall rules",r"Get-NetFirewallRule -DisplayName 'VITEC Scoreboard*' | Get-NetFirewallPortFilter"],["Application processes",r"Get-Process VITEC.Scoreboard,VITEC.VideoOutput,ffmpeg -ErrorAction SilentlyContinue"],["Application log",r"Get-Content 'C:\ProgramData\VITEC Scoreboard\Logs\VITEC-Scoreboard.log' -Tail 100"],["Installer hash",r"Get-FileHash .\VITEC-Scoreboard-Setup-vX.Y.Z.msi -Algorithm SHA256"]],[2100,7260])

heading("9. Authoritative Security References")
bullets(["Microsoft LocalSystem account: https://learn.microsoft.com/windows/win32/services/localsystem-account", "Microsoft Windows Firewall rule recommendations: https://learn.microsoft.com/windows/security/operating-system-security/network-security/windows-firewall/rules", "Microsoft .NET security advisories: https://github.com/dotnet/announcements/labels/Security", "PostgreSQL security information: https://www.postgresql.org/support/security/", "FFmpeg security information: https://ffmpeg.org/security.html"])
para("This document describes version 0.8.63-era behavior and is an operational security reference, not a substitute for the customer’s security architecture, risk assessment, or incident-response policy.")

OUT.parent.mkdir(parents=True, exist_ok=True); WEB.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT); WEB.write_bytes(OUT.read_bytes())
print(OUT); print(WEB)
